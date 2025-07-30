# tools/build_docs.py
# build_docs.py
from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

from docx import Document
from docx.shared import Pt

import config
from tools.compare_docs import compare_docs, analyze_docs, speaker_set_report, length_stats, \
    guess_speaker_mapping, format_speaker_mapping_report, plot_per_speaker_counts_2
from tools.plots import plot_per_speaker_count, plot_hyp_durations_mapped_from_json, \
    plot_hyp_durations_by_ref_bin_from_json

logger = logging.getLogger(__name__)

# ---------- helpers ----------

_PUNCT_NO_LEADING_SPACE = set(",.!?;:%)]}”’")
_PUNCT_NO_TRAILING_SPACE = set("([{$£€“‘")

def _smart_detokenize(tokens: List[str]) -> str:
    """
    Join tokens into readable text:
      - add spaces between words
      - no space before , . ! ? ; : % ) ] } ” ’
      - no space after ( [ { $ £ € “ ‘
    Also collapses extra spaces and fixes simple quote spacing.
    """
    out: List[str] = []
    prev = ""
    for t in tokens:
        if not t:
            continue
        if not out:
            out.append(t)
        else:
            join_with_space = True
            if t in _PUNCT_NO_LEADING_SPACE:
                join_with_space = False
            if prev and prev[-1] in _PUNCT_NO_TRAILING_SPACE:
                join_with_space = False
            out.append(("" if not join_with_space else " ") + t)
        prev = out[-1]
    text = "".join(out)
    # normalize spaces
    text = re.sub(r"\s+", " ", text).strip()
    # simple quote normalization (optional)
    text = text.replace(" ,", ",").replace(" .", ".").replace(" !", "!").replace(" ?", "?").replace(" ;", ";").replace(" :", ":")
    return text

def _words_from_segment(seg: Dict) -> List[str]:
    """
    Prefer aligned `words` if present; otherwise split `text`.
    """
    if "words" in seg and isinstance(seg["words"], list) and seg["words"]:
        toks: List[str] = []
        for w in seg["words"]:
            # some aligners store {'word': 'Hello', ...}
            val = w.get("word") if isinstance(w, dict) else None
            if isinstance(val, str):
                toks.append(val)
        if toks:
            return toks
    # fallback: tokenise naive text
    raw = seg.get("text", "") or ""
    # Split keeping punctuation tokens separate
    pieces = re.findall(r"\w+|[^\w\s]", raw, flags=re.UNICODE)
    return pieces

def _collapse_turns(segments: List[Dict], speaker_key: str = "speaker") -> List[Tuple[str, str]]:
    """
    Merge consecutive segments from the same speaker into turns.
    Returns list of (speaker_id, text).
    """
    turns: List[Tuple[str, str]] = []
    cur_spk: Optional[str] = None
    cur_tokens: List[str] = []

    for seg in segments:
        spk = seg.get(speaker_key) or "SPEAKER_00"
        toks = _words_from_segment(seg)
        if cur_spk is None:
            cur_spk = spk
            cur_tokens.extend(toks)
        elif spk == cur_spk:
            cur_tokens.extend(toks)
        else:
            # flush
            turns.append((cur_spk, _smart_detokenize(cur_tokens)))
            cur_spk = spk
            cur_tokens = list(toks)
    if cur_spk is not None:
        turns.append((cur_spk, _smart_detokenize(cur_tokens)))
    return turns

def _find_pair(folder: Union[str, Path], prefix: Optional[str], aligned_suffix="_aligned.json", diar_suffix=".diarize.json") -> Tuple[Path, Optional[Path]]:
    """
    Locate aligned and diarization jsons in `folder` (optionally starting with `prefix`).
    """
    folder = Path(folder)
    if prefix:
        aligned = next((p for p in folder.glob(f"{prefix}*{aligned_suffix}") if p.is_file()), None)
        diar = next((p for p in folder.glob(f"{prefix}*{diar_suffix}") if p.is_file()), None)
    else:
        aligned = next((p for p in folder.glob(f"*{aligned_suffix}") if p.is_file()), None)
        diar = next((p for p in folder.glob(f"*{diar_suffix}") if p.is_file()), None)
    if not aligned:
        raise FileNotFoundError(f"No aligned JSON matching '*{aligned_suffix}' found in {folder} (prefix={prefix!r}).")
    return aligned, diar

def _load_json(path: Path) -> Dict:
    return json.loads(path.read_text(encoding="utf-8"))

def _default_title_from_prefix(prefix: Optional[str], aligned_path: Path) -> str:
    if prefix:
        return prefix
    return aligned_path.stem.replace("_aligned", "").replace(".aligned", "")

# ---------- main API ----------

def generate_transcript_docx(
    folder: Union[str, Path],
    prefix: Optional[str] = None,
    *,
    aligned_suffix: str = "_aligned.json",
    diar_suffix: str = ".diarize.json",
    out_name: Optional[str] = None,
    speaker_map: Optional[Dict[str, str]] = None,
    include_header: bool = False,
    title: Optional[str] = None,
    base_font: str = "Times New Roman",
    base_size_pt: int = 12,
) -> Path:
    """
    Build a DOCX like the reference format:
      SpeakerName: flowing text (no timestamps), grouped by turns.

    Args:
      folder: directory containing <prefix>*_aligned.json (+ optionally *.diarize.json)
      prefix: optional filename prefix to disambiguate
      aligned_suffix: suffix for aligned json
      diar_suffix: suffix for diarization json (optional)
      out_name: explicit output filename (.docx). Defaults to "<prefix or stem>.docx"
      speaker_map: optional map {"SPEAKER_00": "General Kriegel", ...}
      include_header: if True, insert a small “English Transcript:” header
      title: optional title paragraph at top (defaults to prefix/stem)
      base_font/base_size_pt: document font defaults

    Returns: Path to the written .docx
    """
    aligned_path, _ = _find_pair(folder, prefix, aligned_suffix, diar_suffix)
    logger.info(f"Generating {aligned_path}")

    data = _load_json(aligned_path)
    segments = data.get("segments") or []
    logger.info(f"Collecting {len(segments)} segments")

    # Build turns
    turns = _collapse_turns(segments, speaker_key="speaker")
    logger.info(f"len(turns): {len(turns)} and turns: {turns}")

    # Create docx
    doc = Document()

    # Title (single line; the reference often shows just dialogue, but allow a title)
    doc_title = title or _default_title_from_prefix(prefix, aligned_path)
    if doc_title:
        p = doc.add_paragraph(doc_title)
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(base_size_pt + 2)
        doc.add_paragraph("")  # blank line

    if include_header:
        hdr = doc.add_paragraph("English Transcript:")
        hdr.runs[0].bold = True
        doc.add_paragraph("")

    # Body: "Speaker: text"
    for spk, text in turns:
        shown = speaker_map.get(spk, spk) if speaker_map else spk
        para = doc.add_paragraph(f"{shown}: {text}")

    # Set style defaults (basic)
    style = doc.styles["Normal"]
    style.font.name = base_font
    style.font.size = Pt(base_size_pt)

    # Output name
    out = Path(folder) / (out_name or f"{_default_title_from_prefix(prefix, aligned_path)}.docx")
    logger.info(f"Writing.... {out}")
    doc.save(str(out))
    logger.info(f"Written finished! {out}")
    return out


if __name__ == "__main__":
    out = generate_transcript_docx(folder=config.JSON_ASR_OUTPUT_DIR, prefix="Aridia Conference Call")
    ref_path = "/Users/pouya/PycharmProjects/SmartSummary/Data/Training/English/Aridia Conference Call/Aridia Conference Transcript.docx"
    diar_json_path = "/Users/pouya/PycharmProjects/SmartSummary/tools/asr_outputs/Aridia Conference Call.diarize.json"
    hyp_path = out

    # 1) High-level one-shot
    report = analyze_docs(ref_path, hyp_path)
    print(report["wer_as_is"])
    print(report["wer_no_speakers"])
    print(report["speakers"])
    print(report["lengths"])

    # 2) Individual calls
    wer_as_is = compare_docs(ref_path, hyp_path, exclude_speaker_names=False)
    wer_no_spk = compare_docs(ref_path, hyp_path, exclude_speaker_names=True)
    spk_report = speaker_set_report(ref_path, hyp_path)
    lens = length_stats(ref_path, hyp_path)

    # 3) Optional plot
    plot_per_speaker_counts_2(ref_path, hyp_path, save_path="per_speaker_counts.png")

    res = guess_speaker_mapping(
        ref_docx=Path(ref_path),
        hyp_docx=Path(hyp_path),
        min_score=0.15,  # tweak if you want stricter/looser matches
    )
    print(format_speaker_mapping_report(res))
    # You can also read res["mapping"] as a dict hyp→{ref, score}

    res = guess_speaker_mapping(Path(ref_path), Path(hyp_path), min_score=0.35)
    plot_per_speaker_count(
        ref_turns=res["ref_turns"],
        hyp_turns=res["hyp_turns"],
        mapping_hyp_to_ref=res["mapping_hyp_to_ref"],
    )

    # 1) Per HYP speaker (labels show mapping “HYP → REF”)
    plot_hyp_durations_mapped_from_json(
        diar_json=diar_json_path,
        mapping_hyp_to_ref=res["mapping_hyp_to_ref"],
        title="HYP speaker durations (mapped to REF labels)",
        save_path="hyp_durations_by_speaker.png",
        show=True,
        min_seconds=1.0,  # optional filter
    )

    # 2) Aggregated into REF bins (all HYP mapped into the same REF are summed)
    plot_hyp_durations_by_ref_bin_from_json(
        diar_json=diar_json_path,
        mapping_hyp_to_ref=res["mapping_hyp_to_ref"],
        title="HYP durations aggregated into REF speakers",
        save_path="hyp_durations_by_ref_bins.png",
        show=True,
    )





# compare_docs.py
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any
import re
import collections
import logging

import docx  # pip install python-docx
import jiwer  # pip install jiwer


# Optional plotting (only if you call plot_per_speaker_counts)
try:
    import matplotlib.pyplot as plt  # pip install matplotlib (optional)
except Exception:
    plt = None


# ----------------------------
# Data structures
# ----------------------------

@dataclass
class Segment:
    speaker: Optional[str]
    text: str


@dataclass
class Transcript:
    segments: List[Segment]

    @property
    def speakers(self) -> List[str]:
        """Unique speaker names in order of first appearance (excluding None)."""
        seen = []
        for s in self.segments:
            if s.speaker and s.speaker not in seen:
                seen.append(s.speaker)
        return seen

    @property
    def full_text(self) -> str:
        """Concatenate all segment texts with spaces (ignores speaker names)."""
        return " ".join(seg.text.strip() for seg in self.segments if seg.text.strip())

    def text_with_speaker_lines(self) -> str:
        """Join as 'SPEAKER: text' lines for inspection/debug."""
        lines = []
        for seg in self.segments:
            if seg.speaker:
                lines.append(f"{seg.speaker}: {seg.text.strip()}")
            else:
                lines.append(seg.text.strip())
        return "\n".join(l for l in lines if l)

    def per_speaker_word_counts(self) -> Dict[str, int]:
        counts = collections.Counter()
        for seg in self.segments:
            spk = seg.speaker or "UNKNOWN"
            counts[spk] += _count_words(seg.text)
        return dict(counts)

# ----------------------------
# Public API
# ----------------------------

def parse_docx_transcript(path: str | Path,
                          language_hint: Optional[str] = None) -> Transcript:
    """
    Parse a .docx transcript that may be:
      - Your REF format with markers (English Transcript: ...).
      - Dialogue with 'SPEAKER: text' lines.
      - Generated doc: headings for speaker, paragraphs for utterances.

    Returns Transcript(segments=[Segment(speaker, text), ...]).
    """
    path = Path(path)
    doc = docx.Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    # 1) Try REF marker block (English Transcript -> next marker)
    block = _extract_ref_block("\n".join(paragraphs))
    if block:
        return _parse_dialogue_lines(block)

    # 2) Try dialogue lines with "SPEAKER: text"
    dlg = _collect_dialogue_from_paragraphs(paragraphs)
    if dlg:
        return Transcript(segments=dlg)

    # 3) Try heading-based speaker sections (common for generated docs)
    by_head = _parse_heading_speaker_sections(doc)
    if by_head:
        return Transcript(segments=by_head)

    # 4) Fallback: treat the whole document as single speaker None
    text = " ".join(p for p in paragraphs if p)
    return Transcript(segments=[Segment(None, text)])


def compare_docs(ref_docx: str | Path,
                 hyp_docx: str | Path,
                 exclude_speaker_names: bool = False) -> Dict[str, Any]:
    """
    Compute WER between two .docx transcripts.
    - exclude_speaker_names=False → compare as-is (speaker names are part of text if present).
    - exclude_speaker_names=True  → strip speaker labels before comparing.

    Returns:
      {
        'wer': float,
        'subs': int, 'dels': int, 'ins': int,
        'ref_words': int, 'hyp_words': int
      }
    """
    ref_t = parse_docx_transcript(ref_docx)
    hyp_t = parse_docx_transcript(hyp_docx)

    if exclude_speaker_names:
        ref_text = _strip_speaker_labels(ref_t.text_with_speaker_lines())
        hyp_text = _strip_speaker_labels(hyp_t.text_with_speaker_lines())
    else:
        ref_text = ref_t.full_text
        hyp_text = hyp_t.full_text

    metrics = _compute_wer_metrics(hyp_text, ref_text)
    metrics["ref_words"] = _count_words(ref_text)
    metrics["hyp_words"] = _count_words(hyp_text)
    return metrics


def speaker_set_report(ref_docx: str | Path,
                       hyp_docx: str | Path) -> Dict[str, Any]:
    """
    Report whether the two docs have the same speakers (by name), and what differs.

    Returns:
      {
        'ref_speakers': [...],
        'hyp_speakers': [...],
        'ref_count': int,
        'hyp_count': int,
        'common': [...],
        'only_in_ref': [...],
        'only_in_hyp': [...]
      }
    """
    ref_t = parse_docx_transcript(ref_docx)
    hyp_t = parse_docx_transcript(hyp_docx)

    ref_set = set(_norm_spk(s) for s in ref_t.speakers)
    hyp_set = set(_norm_spk(s) for s in hyp_t.speakers)

    def _orig_names(names: set[str], original_order: List[str]) -> List[str]:
        o = []
        for s in original_order:
            if _norm_spk(s) in names and s not in o:
                o.append(s)
        return o

    return {
        "ref_speakers": ref_t.speakers,
        "hyp_speakers": hyp_t.speakers,
        "ref_count": len(ref_set),
        "hyp_count": len(hyp_set),
        "common": _orig_names(ref_set & hyp_set, ref_t.speakers + hyp_t.speakers),
        "only_in_ref": _orig_names(ref_set - hyp_set, ref_t.speakers),
        "only_in_hyp": _orig_names(hyp_set - ref_set, hyp_t.speakers),
    }


def length_stats(ref_docx: str | Path,
                 hyp_docx: str | Path) -> Dict[str, Any]:
    """
    Return useful length/coverage stats for both docs:
      - totals (segments, words, chars)
      - per-speaker word counts
      - type-token ratio (TTR)
      - avg words/segment
    """
    ref_t = parse_docx_transcript(ref_docx)
    hyp_t = parse_docx_transcript(hyp_docx)

    def _stats(t: Transcript) -> Dict[str, Any]:
        words = _count_words(t.full_text)
        chars = len(_normalize_ws(t.full_text))
        segs = len(t.segments)
        ttr = _type_token_ratio(t.full_text)
        per_spk = t.per_speaker_word_counts()
        return {
            "segments": segs,
            "words": words,
            "chars": chars,
            "avg_words_per_segment": (words / segs) if segs else 0.0,
            "type_token_ratio": ttr,
            "per_speaker_word_counts": per_spk,
        }

    return {
        "ref": _stats(ref_t),
        "hyp": _stats(hyp_t),
    }


def analyze_docs(ref_docx: str | Path,
                 hyp_docx: str | Path) -> Dict[str, Any]:
    """
    One-shot, higher-level analyzer:
      - WER as-is
      - WER excluding speaker labels
      - Speaker set report
      - Length stats

    Returns a dictionary of all above results.
    """
    return {
        "wer_as_is": compare_docs(ref_docx, hyp_docx, exclude_speaker_names=False),
        "wer_no_speakers": compare_docs(ref_docx, hyp_docx, exclude_speaker_names=True),
        "speakers": speaker_set_report(ref_docx, hyp_docx),
        "lengths": length_stats(ref_docx, hyp_docx),
    }


def plot_per_speaker_counts_2(ref_docx: str | Path,
                            hyp_docx: str | Path,
                            save_path: Optional[str | Path] = None) -> None:
    """
    Optional: bar chart of per-speaker word counts for ref vs hyp.
    Requires matplotlib. Does nothing if matplotlib is unavailable.
    """
    if plt is None:
        logging.warning("matplotlib not available; skipping plot.")
        return

    ref_t = parse_docx_transcript(ref_docx)
    hyp_t = parse_docx_transcript(hyp_docx)

    ref_counts = ref_t.per_speaker_word_counts()
    hyp_counts = hyp_t.per_speaker_word_counts()
    # Use union of speakers for consistent ordering
    speakers = sorted(set(list(ref_counts.keys()) + list(hyp_counts.keys())),
                      key=lambda s: (s != "UNKNOWN", s))

    ref_vals = [ref_counts.get(s, 0) for s in speakers]
    hyp_vals = [hyp_counts.get(s, 0) for s in speakers]

    x = range(len(speakers))
    width = 0.4

    plt.figure()
    plt.bar([i - width/2 for i in x], ref_vals, width=width, label="REF")
    plt.bar([i + width/2 for i in x], hyp_vals, width=width, label="HYP")
    plt.xticks(list(x), speakers, rotation=45, ha="right")
    plt.ylabel("Word count")
    plt.title("Per-speaker word counts (REF vs HYP)")
    plt.tight_layout()
    plt.legend()

    if save_path:
        plt.savefig(str(save_path), dpi=150)
    else:
        plt.show()


# ----------------------------
# Internals
# ----------------------------

REF_PATTERNS = [
    r"English\s*Transcript\s*:\s*(.*?)\s*(?=Hebrew\s*Translation\s*:)",
    r"Original\s*Text\s*:\s*(.*?)\s*(?=Translated\s*Text\s*:)",
    r"Source\s*Text\s*:\s*(.*?)\s*(?=Translation\s*:)",
    r"(?:English Transcript|Original Text)\s*:\s*(.*?)\s*(?:Hebrew Translation|Translated Text)\s*:",
]

SPEAKER_LINE_RE = re.compile(r"^\s*([A-Z][A-Za-z0-9 _\-./]{0,64})\s*[:：]\s*(.+)$")


def _extract_ref_block(full_text: str) -> Optional[str]:
    """Find REF block between known markers."""
    for pat in REF_PATTERNS:
        m = re.search(pat, full_text, flags=re.IGNORECASE | re.DOTALL)
        if m:
            return " ".join(m.group(1).split())
    return None


def _collect_dialogue_from_paragraphs(paragraphs: List[str]) -> List[Segment]:
    """
    Collect lines that look like 'SPEAKER: text' from paragraphs.
    """
    segs: List[Segment] = []
    for p in paragraphs:
        if not p:
            continue
        m = SPEAKER_LINE_RE.match(p)
        if m:
            spk = m.group(1).strip()
            txt = m.group(2).strip()
            if txt:
                segs.append(Segment(spk, txt))
    return segs


def _parse_heading_speaker_sections(doc: docx.document.Document) -> List[Segment]:
    """
    Parse documents where speaker names are in headings and the following
    paragraphs are their utterances until the next heading.
    """
    segs: List[Segment] = []
    current_speaker: Optional[str] = None
    buffer: List[str] = []

    def flush():
        nonlocal buffer, current_speaker, segs
        if buffer:
            text = " ".join(t.strip() for t in buffer if t.strip())
            if text:
                segs.append(Segment(current_speaker, text))
        buffer = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        style = (p.style.name or "").lower()
        is_heading = style.startswith("heading")
        looks_like_speaker = bool(SPEAKER_LINE_RE.match(txt))

        if is_heading:
            # new speaker section
            flush()
            # treat raw heading as speaker name (strip any trailing ':')
            current_speaker = txt.rstrip(":：").strip()
        elif looks_like_speaker:
            # explicit "SPEAKER: text"
            flush()
            m = SPEAKER_LINE_RE.match(txt)
            current_speaker = m.group(1).strip() if m else current_speaker
            post = m.group(2).strip() if m else txt
            buffer.append(post)
            flush()
        else:
            buffer.append(txt)

    flush()
    return segs


def _strip_speaker_labels(text: str) -> str:
    """
    Remove leading 'SPEAKER: ' labels from lines. Keeps the spoken text.
    """
    out_lines = []
    for line in text.splitlines():
        m = SPEAKER_LINE_RE.match(line)
        out_lines.append(m.group(2).strip() if m else line.strip())
    return "\n".join(l for l in out_lines if l)

def _collapse_spaces(text: str) -> str:
    """Normalize whitespace without lowercasing/punctuation stripping."""
    # Convert all kinds of whitespace -> spaces, then squeeze.
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    # Normalize multiple newlines -> single space (WER is word-based anyway).
    text = re.sub(r"\s*\n\s*", " ", text)
    return text.strip()


def _compute_wer_metrics(hyp: str, ref: str) -> Dict[str, Any]:
    """
    Compute WER and counts using jiwer with minimal normalization (whitespace only).
    """
    ref_clean = _collapse_spaces(ref)
    hyp_clean = _collapse_spaces(hyp)

    # jiwer
    wer_value = jiwer.wer(ref_clean, hyp_clean)
    measures = jiwer.process_characters(ref_clean, hyp_clean)
    return {
        "wer": wer_value,
        "substitutions": measures.substitutions,
        "deletions": measures.deletions,
        "insertions": measures.insertions,
        # "measures": measures,
    }

def _count_words(s: str) -> int:
    return len(_normalize_ws(s).split())


def _normalize_ws(s: str) -> str:
    return " ".join(s.split())


def _type_token_ratio(s: str) -> float:
    words = [w for w in _normalize_ws(s).lower().split() if w]
    if not words:
        return 0.0
    return len(set(words)) / len(words)


def _norm_spk(name: Optional[str]) -> str:
    if not name:
        return "unknown"
    return re.sub(r"\s+", " ", name.strip().lower())

# --- Speaker mapping utilities -------------------------------------------------
import re
import string
from collections import defaultdict
from pathlib import Path
from typing import Dict, List, Tuple, Optional
import docx

def _normalize_text_for_match(s: str) -> str:
    # Lowercase, strip punctuation, collapse whitespace
    table = str.maketrans({c: " " for c in string.punctuation + "“”‘’—–‑"})
    s = s.lower().translate(table)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def _jaccard(a_tokens: List[str], b_tokens: List[str]) -> float:
    A, B = set(a_tokens), set(b_tokens)
    if not A or not B:
        return 0.0
    return len(A & B) / len(A | B)

def _seq_ratio(a: str, b: str) -> float:
    # Lightweight similarity without extra deps
    import difflib
    return difflib.SequenceMatcher(None, a, b).ratio()

def _similarity_score(a: str, b: str) -> float:
    """
    Combine sequence ratio and Jaccard on word sets.
    Tuned to be robust when content is short or repeated.
    """
    a_n = _normalize_text_for_match(a)
    b_n = _normalize_text_for_match(b)
    j = _jaccard(a_n.split(), b_n.split())
    r = _seq_ratio(a_n, b_n)
    return 0.6 * r + 0.4 * j

def _parse_dialogue_docx(doc_path: Path) -> List[Tuple[str, str]]:
    """
    Parse a dialogue-style DOCX into [(speaker, text), ...].

    Heuristics:
    - If a paragraph contains a colon in the first ~25 chars, treat the prefix as a speaker tag.
    - Otherwise, treat paragraph as a continuation of the last speaker (or 'NARRATOR' if none).
    """
    doc = docx.Document(str(doc_path))
    items: List[Tuple[str, str]] = []
    current_speaker = "NARRATOR"
    buffer = []

    def flush():
        nonlocal buffer, current_speaker
        if buffer:
            txt = " ".join(t.strip() for t in buffer if t.strip())
            if txt:
                items.append((current_speaker, txt))
            buffer = []

    for p in doc.paragraphs:
        line = p.text.strip()
        if not line:
            continue

        # Speaker pattern: "<name>: text"
        # Limit speaker label length to avoid false positives.
        m = re.match(r"^(.{1,25}?):\s*(.+)$", line)
        if m:
            flush()
            spk = m.group(1).strip()
            txt = m.group(2).strip()
            # Normalize very common forms like "Speaker 1", "SPEAKER_00", etc.
            spk_norm = re.sub(r"\s+", "_", spk.upper())
            current_speaker = spk_norm
            buffer.append(txt)
        else:
            buffer.append(line)

    flush()
    return items

def _speaker_profiles(turns: List[Tuple[str, str]]) -> Dict[str, Dict[str, object]]:
    """
    Aggregate dialogue turns into per-speaker profiles:
      {speaker: {"text": concatenated_text, "count": n_turns, "chars": total_len}}
    """
    prof: Dict[str, Dict[str, object]] = defaultdict(lambda: {"text": "", "count": 0, "chars": 0})
    for spk, txt in turns:
        prof[spk]["text"] = (prof[spk]["text"] + " " + txt).strip()
        prof[spk]["count"] = int(prof[spk]["count"]) + 1
        prof[spk]["chars"] = int(prof[spk]["chars"]) + len(txt)
    return dict(prof)
#
# def guess_speaker_mapping(
#     ref_docx: Path,
#     hyp_docx: Path,
#     min_score: float = 0.35
# ) -> Dict[str, Dict[str, object]]:
#     """
#     Best-guess 1→1 mapping from hypothesis speakers to reference speakers.
#
#     Returns:
#       {
#         "mapping": { "HYP_SPEAKER_X": {"ref": "REF_SPEAKER_Y", "score": 0.82}, ...},
#         "unmatched_ref": ["REF_SPK", ...],
#         "unmatched_hyp": ["HYP_SPK", ...],
#         "matrix": [  # optional scores table for inspection
#             {"hyp": "HYP_SPK", "ref": "REF_SPK", "score": 0.xx}, ...
#         ]
#       }
#     """
#     ref_turns = _parse_dialogue_docx(Path(ref_docx))
#     hyp_turns = _parse_dialogue_docx(Path(hyp_docx))
#
#     ref_prof = _speaker_profiles(ref_turns)
#     hyp_prof = _speaker_profiles(hyp_turns)
#
#     ref_speakers = list(ref_prof.keys())
#     hyp_speakers = list(hyp_prof.keys())
#
#     # Score matrix hyp x ref
#     scores: Dict[Tuple[str, str], float] = {}
#     for h in hyp_speakers:
#         for r in ref_speakers:
#             scores[(h, r)] = _similarity_score(ref_prof[r]["text"], hyp_prof[h]["text"])
#
#     # Greedy 1-1 matching: iterate hyp speakers by descending content length
#     used_ref = set()
#     mapping: Dict[str, Dict[str, object]] = {}
#     hyp_sorted = sorted(hyp_speakers, key=lambda s: int(hyp_prof[s]["chars"]), reverse=True)
#     for h in hyp_sorted:
#         # pick best available ref
#         best_r, best_score = None, -1.0
#         for r in ref_speakers:
#             if r in used_ref:
#                 continue
#             sc = scores[(h, r)]
#             if sc > best_score:
#                 best_r, best_score = r, sc
#         if best_r is not None and best_score >= min_score:
#             mapping[h] = {"ref": best_r, "score": round(best_score, 3)}
#             used_ref.add(best_r)
#
#     unmatched_ref = [r for r in ref_speakers if r not in used_ref]
#     unmatched_hyp = [h for h in hyp_speakers if h not in mapping]
#
#     matrix_report = [
#         {"hyp": h, "ref": r, "score": round(scores[(h, r)], 3)}
#         for h in hyp_speakers for r in ref_speakers
#     ]
#     return {
#         "mapping": mapping,
#         "unmatched_ref": unmatched_ref,
#         "unmatched_hyp": unmatched_hyp,
#         "matrix": matrix_report,
#     }

def guess_speaker_mapping(
    ref_docx: Path,
    hyp_docx: Path,
    min_score: float = 0.35
) -> Dict[str, Any]:
    """
    Best-guess 1→1 mapping from hypothesis speakers to reference speakers.

    Returns (backward-compatible + extras):
      {
        # (OLD) hyp->ref mapping, same as before
        "mapping": { "HYP_SPEAKER_X": {"ref": "REF_SPEAKER_Y", "score": 0.82}, ...},

        # (OLD)
        "unmatched_ref": ["REF_SPK", ...],
        "unmatched_hyp": ["HYP_SPK", ...],
        "matrix": [ {"hyp": "HYP_SPK", "ref": "REF_SPK", "score": 0.xx}, ... ],

        # (NEW) for downstream plots/analytics
        "mapping_hyp_to_ref": { "HYP_SPK": "REF_SPK", ... },   # plain map for convenience
        "ref_turns": [...],   # whatever _parse_dialogue_docx returns
        "hyp_turns": [...],   # whatever _parse_dialogue_docx returns
      }
    """
    # Parse turns
    ref_turns = _parse_dialogue_docx(Path(ref_docx))
    hyp_turns = _parse_dialogue_docx(Path(hyp_docx))

    # Build simple speaker profiles (e.g., concatenated text per speaker)
    ref_prof = _speaker_profiles(ref_turns)
    hyp_prof = _speaker_profiles(hyp_turns)

    ref_speakers = list(ref_prof.keys())
    hyp_speakers = list(hyp_prof.keys())

    # Score matrix hyp x ref
    scores: Dict[Tuple[str, str], float] = {}
    for h in hyp_speakers:
        for r in ref_speakers:
            scores[(h, r)] = _similarity_score(ref_prof[r]["text"], hyp_prof[h]["text"])

    # Greedy 1-1 matching: iterate hyp speakers by descending content length
    used_ref = set()
    mapping: Dict[str, Dict[str, Any]] = {}
    hyp_sorted = sorted(hyp_speakers, key=lambda s: int(hyp_prof[s]["chars"]), reverse=True)
    for h in hyp_sorted:
        best_r, best_score = None, -1.0
        for r in ref_speakers:
            if r in used_ref:
                continue
            sc = scores[(h, r)]
            if sc > best_score:
                best_r, best_score = r, sc
        if best_r is not None and best_score >= min_score:
            mapping[h] = {"ref": best_r, "score": round(best_score, 3)}
            used_ref.add(best_r)

    unmatched_ref = [r for r in ref_speakers if r not in used_ref]
    unmatched_hyp = [h for h in hyp_speakers if h not in mapping]

    matrix_report = [
        {"hyp": h, "ref": r, "score": round(scores[(h, r)], 3)}
        for h in hyp_speakers for r in ref_speakers
    ]

    # NEW: plain hyp->ref mapping for plotting
    mapping_hyp_to_ref: Dict[str, str] = {h: v["ref"] for h, v in mapping.items()}

    return {
        "mapping": mapping,                     # original structure (with scores)
        "unmatched_ref": unmatched_ref,
        "unmatched_hyp": unmatched_hyp,
        "matrix": matrix_report,

        # extras for plots/analytics
        "mapping_hyp_to_ref": mapping_hyp_to_ref,
        "ref_turns": ref_turns,
        "hyp_turns": hyp_turns,
    }

def format_speaker_mapping_report(result: Dict[str, Dict[str, object]]) -> str:
    """
    Pretty-print the mapping guess:
    """
    lines = []
    lines.append("Best‑guess speaker mapping (hyp → ref):")
    if not result["mapping"]:
        lines.append("  (no confident mappings)")
    else:
        # Show in order of decreasing confidence
        items = sorted(result["mapping"].items(), key=lambda kv: kv[1]["score"], reverse=True)
        for hyp, info in items:
            lines.append(f"  {hyp} → {info['ref']}  (score={info['score']})")

    if result["unmatched_hyp"]:
        lines.append("\nUnmatched HYP speakers:")
        for h in result["unmatched_hyp"]:
            lines.append(f"  - {h}")

    if result["unmatched_ref"]:
        lines.append("\nUnmatched REF speakers:")
        for r in result["unmatched_ref"]:
            lines.append(f"  - {r}")

    return "\n".join(lines)

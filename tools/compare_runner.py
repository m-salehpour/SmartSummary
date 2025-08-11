# tools/compare_runner.py
from __future__ import annotations

import difflib
import json
import logging
import re
from pathlib import Path
from typing import Dict, Any, Optional, List, Union, Sequence, Tuple

import config

try:
    import docx  # python-docx
except Exception:
    docx = None  # allow import even if docx isn't installed

import jiwer

from tools.normalizers import _run_llm_clean
logger = logging.getLogger(__name__)



def run_basic_comparisons(
    raw_segments: list[dict],
    ref_path: Path,
    lang: Optional[str],
    diff: bool,
    print_hyp: bool,
    print_ref: bool,
    fallback: str = config.FALLBACK_POLICY_FULL,
    strip_speakers: bool =True,
    script_hint: str ="latin",
    prefix_message: str ="",
) -> Dict[str, Any]:
    """
    1) Build hyp from raw segments
    2) Load REF once
    3) Compare RAW and NO-LLM-cleaned

    Returns:
      {
        "ref_text": str,
        "hyp_raw": str,
        "hyp_no_llm": str,
        "metrics_raw": {...},
        "metrics_no_llm": {...}
      }
    """
    # REF (once)
    ref_text = load_reference_text(
        str(ref_path),
        fallback=fallback,
        strip_speakers=strip_speakers,
        script_hint=script_hint,
    ).strip()

    if print_ref:
        print("\n[REF TEXT]\n", ref_text)

    # RAW hyp
    hyp_raw = get_hypothesis_text(raw_segments).strip()
    if print_hyp:
        print("\n[HYP RAW]\n", hyp_raw)
    print("\n=== RAW TRANSCRIPTION ===")
    metrics_raw = compare_texts(hyp_raw, ref_text, diff=diff, prefix_message=prefix_message)

    # NO-LLM hyp
    # Defer to your language-aware normalizer that you import in ASR and pass here if desired
    # (We keep the interface clean – the caller can provide a callable or reuse the one in asr)
    return {
        "ref_text": ref_text,
        "hyp_raw": hyp_raw,
        "metrics_raw": metrics_raw,
    }


def run_no_llm_clean_and_compare(
    hyp_no_llm: str,
    ref_text: str,
    diff: bool,
    print_hyp: bool,
    prefix_message: str ="",
) -> Dict[str, Any]:
    if print_hyp:
        print("\n[HYP NO-LLM CLEAN]\n", hyp_no_llm)
    print("\n=== NO-LLM CLEANED ===")
    metrics_no_llm = compare_texts(hyp_no_llm, ref_text, diff=diff, prefix_message=prefix_message)
    return {
        "hyp_no_llm": hyp_no_llm,
        "metrics_no_llm": metrics_no_llm,
    }


def run_llm_clean_from_raw(
    raw_json_path: Path,
    ref_text: str,
    diff: bool,
    print_hyp: bool,
    suffix: str = "_llm_from_raw",
    prefix_message: str ="",
) -> Dict[str, Any]:
    """
    Runs LLM cleaner on the RAW transcript JSON saved by ASR.
    """
    llm_raw_json = _run_llm_clean(raw_json_path, suffix=suffix)
    llm_raw_data = json.loads(llm_raw_json.read_text(encoding="utf-8"))
    hyp_llm_raw  = get_hypothesis_text(llm_raw_data["segments"]).strip()
    if print_hyp:
        print("\n[HYP LLM(from raw)]\n", hyp_llm_raw)
    print("\n=== LLM-CLEANED FROM RAW ===")
    metrics_llm_raw = compare_texts(hyp_llm_raw, ref_text, diff=diff, prefix_message=prefix_message)
    return {
        "hyp_llm_from_raw": hyp_llm_raw,
        "metrics_llm_from_raw": metrics_llm_raw,
        "llm_raw_json": str(llm_raw_json),
    }


def run_llm_clean_from_no_llm(
    raw_json_path: Path,
    hyp_no_llm: str,
    ref_text: str,
    diff: bool,
    print_hyp: bool,
    suffix_intermediate: str = "_no_llm_cleaned",
    suffix_final: str = "_llm_from_no_llm",
    prefix_message: str ="",
) -> Dict[str, Any]:
    """
    Writes a minimal JSON for the NO-LLM string, runs LLM cleaner on it, then compares.
    """
    no_llm_json = raw_json_path.with_name(
        raw_json_path.stem + suffix_intermediate + raw_json_path.suffix
    )
    with open(no_llm_json, "w", encoding="utf-8") as f:
        json.dump({"segments": [{"text": hyp_no_llm}]}, f, indent=2)

    llm_no_llm_json = _run_llm_clean(no_llm_json, suffix=suffix_final)
    llm_no_llm_data = json.loads(llm_no_llm_json.read_text(encoding="utf-8"))
    hyp_llm_no_llm  = get_hypothesis_text(llm_no_llm_data["segments"]).strip()
    if print_hyp:
        print("\n[HYP LLM(from no-LLM)]\n", hyp_llm_no_llm)
    print("\n=== LLM-CLEANED FROM NO-LLM ===")
    metrics_llm_no_llm = compare_texts(hyp_llm_no_llm, ref_text, diff=diff, prefix_message=prefix_message)
    return {
        "hyp_llm_from_no_llm": hyp_llm_no_llm,
        "metrics_llm_from_no_llm": metrics_llm_no_llm,
        "llm_no_llm_json": str(llm_no_llm_json),
    }




def get_hypothesis_text(segments: List[Dict[str, Any]]) -> str:
    """
    Concatenate ASR segments into a single string.
    """
    return " ".join(seg.get("text", "").strip() for seg in segments)


# def load_reference_text(path: Union[str, Path]) -> str:
#     """
#     Read a .docx file and extract the reference transcript between known markers.
#     Raises ValueError if markers not found.
#     """
#     path = Path(path)
#     doc = docx.Document(str(path))
#     full_text = "\n".join(p.text for p in doc.paragraphs)
#     # Define start/end marker patterns
#     patterns = [
#         r"English\s*Transcript\s*:\s*(.*?)\s*(?=Hebrew\s*Translation\s*:)",
#         r"Original\s*Text\s*:\s*(.*?)\s*(?=Translated\s*Text\s*:)",
#         r"Source\s*Text\s*:\s*(.*?)\s*(?=Translation\s*:)",
#         r"(?:English Transcript|Original Text)\s*:\s*(.*?)\s*(?:Hebrew Translation|Translated Text)\s*:",
#     ]
#     for pat in patterns:
#         match = re.search(pat, full_text, flags=re.IGNORECASE | re.DOTALL)
#         if match:
#             raw = match.group(1)
#             return " ".join(raw.split())
#     # No match found, raise with snippet
#     snippet = full_text[:500].replace("\n", " ")
#     raise ValueError(
#         f"Could not find expected markers in '{path}'. Snippet: {snippet}"
#     )



DEFAULT_MARKERS: Tuple[str, ...] = (
    r"English\s*Transcript\s*:\s*(.*?)\s*(?=Hebrew\s*Translation\s*:)",
    r"Original\s*Text\s*:\s*(.*?)\s*(?=Translated\s*Text\s*:)",
    r"Source\s*Text\s*:\s*(.*?)\s*(?=Translation\s*:)",
    r"(?:English Transcript|Original Text)\s*:\s*(.*?)\s*(?:Hebrew Translation|Translated Text)\s*:",
)

# Simple Unicode script ranges (coarse heuristic)
SCRIPT_RANGES = {
    "latin": r"\u0000-\u024F",
    "hebrew": r"\u0590-\u05FF",
    "arabic": r"\u0600-\u06FF",   # covers Arabic + Persian block
}

def _collapse_ws(s: str) -> str:
    return " ".join(s.split())

def _read_docx(path: Path) -> str:
    if docx is None:
        raise RuntimeError("python-docx not installed; cannot read .docx")
    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)

def _read_text(path: Path) -> str:
    return Path(path).read_text(encoding="utf-8", errors="replace")

def _extract_between_markers(text: str, patterns: Sequence[str]) -> Optional[str]:
    for pat in patterns:
        m = re.search(pat, text, flags=re.IGNORECASE | re.DOTALL)
        if m:
            return _collapse_ws(m.group(1))
    return None

def _strip_speaker_labels(text: str) -> str:
    """
    Convert lines like 'General Kriegel: We are ready.' -> 'We are ready.'
    Only strips when a clear 'Name: ' prefix exists.
    """
    lines = []
    for line in text.splitlines():
        # remove 'Something:' at start (allow unicode letters and spaces)
        stripped = re.sub(r"^\s*[\w .,\-\u0590-\u05FF\u0600-\u06FF]+:\s+", "", line)
        lines.append(stripped)
    return "\n".join(lines)

try:
    import regex as _re  # supports \p{...}
    _HAS_REGEX = True
except Exception:
    import re as _re
    _HAS_REGEX = False

import unicodedata

# Map script hints to Unicode script properties (regex module)
_SCRIPT_PROP = {
    "latin":  r"\p{Latin}",
    "hebrew": r"\p{Hebrew}",
    "arabic": r"\p{Arabic}",
    # add others if you need them (Greek, Cyrillic, etc.)
}

def _filter_by_script(text: str, script_hint: str | None) -> str:
    """
    Keep only characters from the hinted script plus punctuation & spaces.
    Falls back to a manual filter when 'regex' is unavailable.
    """
    if not script_hint:
        return text  # no filtering requested

    if _HAS_REGEX and script_hint.lower() in _SCRIPT_PROP:
        # Allow characters in the script, plus punctuation (\p{P}) and spaces (\p{Zs})
        script_pat = _SCRIPT_PROP[script_hint.lower()]
        # Remove any char NOT in (script OR punctuation OR space)
        return _re.sub(fr"(?!{script_pat}|\p{{P}}|\p{{Zs}}).", "", text)

    # Fallback: built-in 're' – approximate via codepoint ranges and unicodedata
    ranges = {
        "latin":  ((0x0041, 0x024F),),
        "hebrew": ((0x0590, 0x05FF), (0xFB1D, 0xFB4F)),  # Hebrew + presentation forms
        "arabic": ((0x0600, 0x06FF), (0x0750, 0x077F), (0x08A0, 0x08FF), (0xFB50, 0xFDFF), (0xFE70, 0xFEFF)),
    }
    allowed_ranges = ranges.get(script_hint.lower())
    if not allowed_ranges:
        return text

    def _ok(ch: str) -> bool:
        cp = ord(ch)
        # keep punctuation and spaces
        cat = unicodedata.category(ch)  # 'P*' = punctuation, 'Zs'=space separator
        if cat.startswith("P") or cat == "Zs":
            return True
        # keep if within any of the allowed script ranges
        return any(lo <= cp <= hi for (lo, hi) in allowed_ranges)

    return "".join(ch for ch in text if _ok(ch))

def load_reference_text(
    path: Union[str, Path],
    *,
    patterns: Sequence[str] = DEFAULT_MARKERS,
    # fallback policy when markers missing:
    #   "full"     -> return whole document (after cleanup)
    #   "dialogue" -> strip "Speaker: " labels, then return
    #   "first"    -> first non-empty block (>= min_chars)
    #   "none"     -> raise ValueError (strict mode)
    fallback: str = config.FALLBACK_POLICY_FULL,
    min_chars: int = 120,
    # optional cleanups
    strip_speakers: bool = False,
    script_hint: Optional[str] = None,  # e.g., "latin", "hebrew", "arabic"
) -> str:
    """
    Robust transcript loader.

    1) If DOCX/TXT has known section markers, return that span.
    2) Otherwise apply a chosen fallback:
        - 'dialogue': strip speaker labels and return all text
        - 'full':     return entire text
        - 'first':    return the first sufficiently long block
        - 'none':     raise ValueError (strict)
    3) Optional: filter to a script (latin/hebrew/arabic) and collapse spaces.

    Examples:
        # Aridia .docx has no markers and uses 'Speaker: text'
        txt = load_reference_text("Aridia Conference Transcript.docx",
                                  fallback="dialogue", strip_speakers=True)

        # Plain Farsi .txt: keep full text, filtered to Arabic/Persian block
        txt = load_reference_text("response_..._farsi.txt",
                                  fallback="full", script_hint="arabic")
    """
    path = Path(path)
    ext = path.suffix.lower()

    # Read raw
    if ext == ".docx":
        raw = _read_docx(path)
    elif ext in {".txt", ".md"}:
        raw = _read_text(path)
    else:
        # Try text read; if it fails and it's not docx, let it raise
        try:
            raw = _read_text(path)
        except Exception:
            if docx and ext in {".doc", ".docm"}:
                raw = _read_docx(path)
            else:
                raise

    # 1) Try markers
    extracted = _extract_between_markers(raw, patterns)
    if extracted:
        text = extracted
    else:
        # 2) Fallback strategy
        if fallback not in {"full", "dialogue", "first", "none"}:
            logger.warning("Unknown fallback '%s'; using 'full'.", fallback)
            fallback = "full"

        logger.info("Reference markers not found in %s; applying fallback='%s'.", path.name, fallback)

        cleaned = raw
        if strip_speakers or fallback == "dialogue":
            cleaned = _strip_speaker_labels(cleaned)

        if fallback == "first":
            blocks = [b.strip() for b in re.split(r"\n{2,}", cleaned)]
            block = next((b for b in blocks if len(_collapse_ws(b)) >= min_chars), "")
            if not block:
                logger.warning("No block >= %d chars found; falling back to 'full'.", min_chars)
                text = cleaned
            else:
                text = block
        elif fallback == "none":
            snippet = _collapse_ws(raw[:300])
            raise ValueError(f"Could not find expected markers in '{path}'. Snippet: {snippet}")
        else:
            # 'full' or 'dialogue'
            text = cleaned

    # 3) Optional script filter
    if script_hint:
        text = _filter_by_script(text, script_hint)

    return _collapse_ws(text)


def compute_wer_metrics(hyp: str, ref: str) -> Dict[str, Any]:
    """
    Compute WER and related counts between reference and hypothesis.
    Returns a dict with 'wer', 'substitutions', 'deletions', 'insertions'.
    """
    wer_value = jiwer.wer(ref, hyp)
    measures = jiwer.process_characters(ref, hyp)
    return {
        "wer": wer_value,
        "substitutions": measures.substitutions,
        "deletions": measures.deletions,
        "insertions": measures.insertions,
    }


def diff_word_level(ref: str, hyp: str) -> List[str]:
    """
    Return a unified diff of reference vs hypothesis word lists.
    """
    ref_words = ref.split()
    hyp_words = hyp.split()
    return list(difflib.unified_diff(
        ref_words,
        hyp_words,
        fromfile="REFERENCE",
        tofile="HYPOTHESIS",
        lineterm=""
    ))


def compare_texts(
    hyp: str,
    ref: str,
    diff: bool = False,
    prefix_message: str = "",
) -> Dict[str, Any]:
    """
    Compute and optionally print WER metrics and word-level diff.
    Returns the metrics dict.
    """
    metrics = compute_wer_metrics(hyp, ref)
    logging.info(f"📊{prefix_message} WER: {metrics['wer']:.2%} | S={metrics['substitutions']} D={metrics['deletions']} I={metrics['insertions']}")
    if diff:
        diffs = diff_word_level(ref, hyp)
        logging.info("🔍 Word-level diff:")
        for line in diffs:
            logging.info(line)
    return metrics


def segments_comparison(
    hyp_text: str,
    ground_truth_path: Union[str, Path],
    audio_path: Union[str, Path],
    diff: bool = False,
    print_hyp: bool = False,
    print_ref: bool = False,
    msg: Optional[str] = None,
    lang: Optional[str] = None,
) -> Dict[str, Dict[str, Any]]:
    """
    Full pipeline: extract hypothesis, load reference, normalize, and compare.

    Returns a dict with 'raw' and 'clean' metrics.
    """
    # Extract texts
    ref_text = load_reference_text(ground_truth_path)

    if print_hyp:
        logging.info(f"[Raw hyp] {hyp_text}")
    if print_ref:
        logging.info(f"[Ref] {ref_text}")

    # Raw comparison
    logging.info(f"--- Comparing raw {msg} ---")
    raw_metrics = compare_texts(hyp_text, ref_text, diff)

    clean_metrics = ""

    return {"raw": raw_metrics, "clean": clean_metrics}


def compare_segments(
    segments,
    ground_truth_path: Path,
    audio_path: Path,
    diff: bool,
    print_hyp: bool,
    print_ref: bool,
    msg: str,
    lang: str,
):
    print(f"\n--- Comparing {msg} ({lang}) ---\n")
    segments_comparison(
        segments,
        str(ground_truth_path),
        str(audio_path),
        diff=diff,
        print_hyp=print_hyp,
        print_ref=print_ref,
        msg=msg,
        lang=lang,
    )


def _compare_strs(hyp: str, ref: str, **cmp_kwargs):
    """Helper to do a single string-based comparison and print WER/diff."""
    metrics = compare_texts(hyp, ref, **cmp_kwargs)
    return metrics
